from docx import Document
import os

# Use absolute path
script_dir = os.path.dirname(os.path.abspath(__file__))
doc_path = os.path.join(script_dir, 'PGHeavenDocumentation.docx')

# Try to open existing document, if it fails, create a new one
try:
    print(f'Attempting to open: {doc_path}')
    doc = Document(doc_path)
    print('Successfully opened existing document')
except Exception as e:
    print(f'Failed to open document, creating new one: {e}')
    doc = Document()
    # Add Chapter 1 content
    chapter1_title = '''Chapter 1: Introduction

1.1 Project Overview

The "PG Heaven" web application is a complete platform designed to bridge the gap between paying‑guest (PG) owners and students, young professionals or travellers seeking affordable, comfortable accommodation.
It consists of a Node.js/Express back‑end serving a MongoDB data store, and an Angular front‑end providing a responsive, modern UI.
Users can register as owners or tenants, browse PG listings, book rooms, lodge complaints or service requests, and make payments. Administrators manage listings, monitor bookings and handle customer care tasks.

The application models real‑world PG management and is intended to:

* mimic an online marketplace for shared housing,
* support both mobile and desktop clients via a SPA,
* facilitate secure authentication, file uploads (room photos), and transactional operations.

1.2 Objectives

The primary goals of the project are:

1. Simplify PG discovery: provide a searchable, filterable catalogue of rooms & hostels across cities.
2. Streamline bookings: allow tenants to reserve rooms on‑line with real‑time availability checks.
3. Enable owner/manager control: permit PG proprietors to list/update rooms, receive bookings and respond to service issues.
4. Enhance user experience: design a clean UI with role‑based dashboards (admin, owner, tenant).
5. Ensure security & stability: implement JWT‑based auth, input validation, and proper error handling.
6. Provide extensibility: structure code so new features (e.g. payment gateways, reviews, analytics) can be added later.

1.3 Scope

This documentation and the corresponding implementation cover the following areas:

* Back‑end services  
  - RESTful APIs for authentication, room management, bookings, complaints, payments, etc.  
  - MongoDB schemas/models (User, PG, Room, Booking, Complaint, ServiceRequest, Order, Invoice).

* Front‑end application  
  - Angular components for browsing, booking, complaint submission, dashboard interfaces, and admin panels.  
  - Services responsible for communicating with the API.

* Middleware and utilities  
  - Authentication middleware, file upload handling, and role‑based guards.  
  - Configuration for database connections and environment settings.

* Data  
  - Sample JSON data files (`pg-rental.rooms.json`, `pg-rental.pgs.json`) included for development/testing.

* Testing and deployment  
  - Basic unit tests (where present) and instructions for running the app locally.  
  - Guidance for packaging and running the Node.js server, as well as building the Angular client.

Out‑of‑scope for this version:

- Integration with external payment gateways (mock payments are used).
- Mobile‑specific apps (only responsive web UI).
- Advanced analytics or machine‑learning features.

1.4 Problem Statement

Finding safe, cost‑effective paying‑guest accommodation is often a time‑consuming process involving multiple calls and visits.
Existing solutions are fragmented, lack transparency, or require significant manual effort to maintain listings. For PG owners, managing bookings, addressing complaints and updating room availability can be cumbersome without a centralized system.

PG Heaven aims to solve these problems by:

* aggregating PG listings in one place,
* providing real-time availability and booking management,
* giving owners an easy-to-use interface for administration,
* and offering tenants the ability to raise issues or request services quickly.

By automating routine tasks and presenting information clearly, the platform reduces friction for both owners and residents while improving overall satisfaction.'''
    
    # Parse and add Chapter 1
    doc.add_heading('Chapter 1: Introduction', level=1)
    for line in chapter1_title.replace('Chapter 1: Introduction\n\n', '').splitlines():
        if line.startswith('1.'):
            doc.add_heading(line, level=2)
        elif line.startswith('*') or line.startswith('-'):
            doc.add_paragraph(line, style='List Bullet')
        elif line.startswith('1.') and '.' in line[1:]:
            continue  # already handled
        else:
            if line.strip():
                doc.add_paragraph(line)

chapter2 = '''Chapter 2: System Analysis

2.1 Existing System

The current process for locating PG accommodations is largely manual. Individuals often rely on word-of-mouth, unorganized listings on social media, or physically visiting properties, which is time-consuming and inefficient. Information about availability, pricing, and amenities is scattered, leading to confusion and frustration.

2.2 Limitations

The existing system suffers from several drawbacks:

- **Scalability issues** – manual listings cannot cope with growing numbers of properties or users.
- **Lack of transparency** – prices and room statuses are not updated in real time, resulting in double bookings or incorrect information.
- **No centralized control** – owners must manage bookings through spreadsheets or paper records.
- **Poor user experience** – tenants cannot easily filter or compare options, leading to a frustrating search.

2.3 Proposed System

The proposed PG Heaven application introduces a centralized, digital platform to manage PG listings and bookings. Key features include:

* Online registration for owners and tenants
* Searchable, filterable catalogue of rooms by city, price, facilities, etc.
* Real-time availability updates and booking workflow.
* Administrative dashboard for owners to add/edit rooms, view bookings and respond to complaints.
* Integrated payment and complaint management modules.

This system aims to streamline operations, improve information accuracy, and provide a convenient user interface.

2.4 Feasibility Study

A feasibility analysis was conducted across four dimensions:

- **Technical feasibility:** The stack (Node.js, Express, MongoDB, Angular) is well-suited to web deployment and the team has experience with these technologies.
- **Economic feasibility:** Development costs are minimal using open-source tools; the system could be monetized via listing fees or premium features.
- **Operational feasibility:** Users are already familiar with web applications, and the system reduces manual workload significantly.
- **Schedule feasibility:** A basic MVP can be delivered within a few months given the modular architecture and reuse of existing templates.

The study concludes that the proposed system is practical and holds significant benefits over current practices.'''

for line in chapter2.splitlines():
    if line.startswith('Chapter 2'):
        doc.add_heading(line, level=1)
    elif line.startswith('2.'):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)

chapter3 = '''Chapter 3: System Design

3.1 System Architecture

The PG Heaven application follows a three-tier architecture consisting of:

* **Presentation Layer (Frontend):** Angular-based Single Page Application (SPA) providing a responsive user interface for different user roles (Admin, Owner, Tenant).
* **Business Logic Layer (API Server):** Node.js with Express framework handling request routing, business logic execution, and data validation.
* **Data Layer (Database):** MongoDB for NoSQL document storage, providing flexibility and scalability for varying data structures.

The system uses RESTful APIs for communication between the frontend and backend, ensuring standard HTTP methods (GET, POST, PUT, DELETE) for CRUD operations.

3.2 Data Flow Diagram (DFD)

The data flow in PG Heaven follows these primary paths:

* User initiates action via Angular frontend
* Frontend sends HTTP request to Express API server with JWT token for authentication
* API server validates request, processes business logic, and queries/updates MongoDB
* MongoDB returns query results to API server
* API server sends JSON response back to frontend
* Frontend updates UI with received data

Key data flows include:
- User authentication and registration
- PG listing browse and search
- Booking room reservation and confirmation
- Payment processing and invoice generation
- Complaint and service request submission and resolution

3.3 Use Case Diagram

The main actors and their use cases include:

* **Admin:** Manage PG listings, view and manage all bookings, handle complaints, monitor payments, and generate reports
* **Owner/PG Manager:** Register PG, add/update rooms, manage bookings, view complaints, update room availability
* **Tenant/Customer:** Register account, search and browse PGs, book rooms, make payments, submit complaints, request services

3.4 Database Design (ER Diagram)

The database schema includes the following entities and relationships:

* **User Entity:** Stores user information including name, email, password, phone number, role, and address.
* **PG Entity:** Represents paying-guest properties with name, location, city, owner reference, and amenities.
* **Room Entity:** Individual rooms within a PG with room number, type, price, capacity, and availability status.
* **Booking Entity:** Records room reservations with booking date, check-in, check-out, status, and tenant reference.
* **Payment Entity:** Tracks payment transactions with amount, payment date, method, and booking reference.
* **Complaint Entity:** Stores customer complaints with description, status, priority, and assigned owner.
* **ServiceRequest Entity:** Manages service requests with type, description, status, and completion date.
* **Invoice Entity:** Generated records of transactions with ID, amount, date, and payment status.

Relationships:
- One User can own multiple PGs (one-to-many)
- One PG contains multiple Rooms (one-to-many)
- One Room has multiple Bookings (one-to-many)
- One Booking generates one Payment (one-to-one)
- One Booking can have multiple Complaints (one-to-many)
- One User can have multiple Complaints (one-to-many)

3.5 UI Design

The user interface is designed with the following principles:

* **Responsive Design:** Supports desktop, tablet, and mobile devices using Bootstrap and CSS Grid.
* **Role-Based Dashboards:** Different views for Admin, Owner, and Tenant with relevant information and actions.
* **Intuitive Navigation:** Clear menu structure and breadcrumbs for easy navigation across pages.
* **Consistent Styling:** Uniform color scheme, typography, and component design across the application.
* **Accessibility:** WCAG 2.1 compliance for keyboard navigation and screen reader support.

Key UI Components:
- Header with navigation and user profile menu
- Sidebar navigation for admin and owner panels
- Search and filter interface for PG browse
- Card-based layout for displaying room listings
- Form components for booking, complaint, and payment
- Modal dialogs for confirmations and alerts
- Data tables with sorting and pagination for management views
- Charts and graphs for analytics dashboards'''

for line in chapter3.splitlines():
    if line.startswith('Chapter 3'):
        doc.add_heading(line, level=1)
    elif line.startswith('3.'):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)

chapter4 = '''Chapter 4: Technology Used

The PG Heaven application leverages a modern, open‑source technology stack:

* **Backend:** Node.js with Express for a lightweight, asynchronous server framework.
* **Database:** MongoDB for flexible, document-oriented data storage.
* **Frontend:** Angular (TypeScript) providing a component-based SPA architecture.
* **Authentication:** JSON Web Tokens (JWT) for stateless, secure API access.
* **Hosting:** Can be deployed on cloud platforms such as Azure App Service, AWS EC2, or Heroku.
* **Miscellaneous:** Mongoose ODM for MongoDB interactions, bcrypt for password hashing, Multer for file uploads, and Bootstrap/SCSS for styling.

The choice of these technologies aims at scalability, developer productivity, and cross-platform compatibility.'''

for line in chapter4.splitlines():
    if line.startswith('Chapter 4'):
        doc.add_heading(line, level=1)
    elif line.startswith('4.'):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)

chapter5 = '''Chapter 5: Implementation

This section outlines key implementation details and file structure:

* **Server setup:** Express server defined in `backend/server.js` with route modules under `controllers` and `routes`.
* **Database models:** Mongoose schemas located in `backend/models` representing User, PG, Room, Booking, Complaint, ServiceRequest, Order, and Invoice.
* **Authentication middleware:** JWT validation implemented in `backend/middleware/auth.js` supporting role-based access control.
* **File uploads:** Handled using Multer configured in `backend/middleware/upload.js` for PG photo uploads.
* **Routing:** RESTful endpoints segmented by function (auth, booking, complaint, menu, order, payment, pg, service) under `backend/routes`.
* **Frontend:** Angular components mirror backend functionality; services (`src/app/services/*.service.ts`) call the API endpoints.
* **Configuration:** Database credentials and environment variables stored in `backend/config/db.js` and Angular `environment.ts`.
* **Scripts:** NPM scripts in root and subfolders manage start, build, and test tasks.

Developers can run `npm install` in both `backend` and `frontend` directories, start the server with `npm run dev` (or `node server.js`), and serve the Angular app with `ng serve` for local development.'''

for line in chapter5.splitlines():
    if line.startswith('Chapter 5'):
        doc.add_heading(line, level=1)
    elif line.startswith('5.'):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)

chapter6 = '''Chapter 6: Testing

Testing was carried out to ensure correctness, security, and performance:

* **Unit tests:** Some backend functions and Angular services include basic Jest/Mocha tests (if present) located under respective `test` folders.
* **Integration tests:** API endpoints were manually verified using Postman collections supplied alongside the project.
* **Frontend testing:** Angular components include simple Karma/Jasmine specs verifying rendering and service interactions.
* **Manual QA:** End-to-end scenarios (registration, booking, payments, complaints) executed in a local environment to validate workflows.

Test results indicated that core features function as expected under normal conditions. Any discovered bugs were noted for future fixes.'''

for line in chapter6.splitlines():
    if line.startswith('Chapter 6'):
        doc.add_heading(line, level=1)
    elif line.startswith('6.'):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)

chapter7 = '''Chapter 7: Results

The implemented system provides a functional PG listing and booking platform. Key outcomes include:

* Successful role-based access and separation of concerns between admins, owners, and tenants.
* Real-time availability updates and avoidance of conflicts during booking tests.
* File upload handling demonstrated with sample PG images.
* Basic payment flow simulated (mock) and invoice records generated.
* Complaint and service request workflows tested end-to-end.

Overall, the MVP meets the project objectives, though certain edge cases and scalability aspects require further development.'''

for line in chapter7.splitlines():
    if line.startswith('Chapter 7'):
        doc.add_heading(line, level=1)
    elif line.startswith('7.'):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)

chapter8 = '''Chapter 8: Conclusion & Future Scope

The PG Heaven platform successfully addresses the need for a centralized PG accommodation marketplace. Conclusion points:

* System facilitates easy discovery and management of PG listings.
* Role-driven dashboards improve operational efficiency for owners and admins.
* Real-world workflows (booking, payment, complaints) are automated.

Future scope includes:

* Integration with real payment gateways and mobile app versions.
* Advanced search filters, reviews and rating modules.
* Analytics dashboard with usage metrics and predictive availability.
* Scalability improvements such as microservices or serverless functions.
* Security enhancements including MFA, rate limiting, and audit logging.'''

for line in chapter8.splitlines():
    if line.startswith('Chapter 8'):
        doc.add_heading(line, level=1)
    elif line.startswith('8.'):
        doc.add_heading(line, level=2)
    else:
        doc.add_paragraph(line)

references = '''References

* Project documentation and sample data provided within the repository.
* Node.js and Express official documentation (https://nodejs.org, https://expressjs.com).
* Angular framework documentation (https://angular.io).
* MongoDB manual (https://docs.mongodb.com).
* Various open-source libraries referenced by package.json files.'''

for line in references.splitlines():
    if line.startswith('References'):
        doc.add_heading(line, level=1)
    else:
        doc.add_paragraph(line)

out_path = os.path.join(script_dir, 'PGHeavenDocumentation.docx')
updated_path = os.path.join(script_dir, 'PGHeavenDocumentation_Updated.docx')

# Save to updated file
doc.save(updated_path)
print('saved to', updated_path)
print('Note: Close PGHeavenDocumentation.docx in your editor/application, then replace it with this updated version')
